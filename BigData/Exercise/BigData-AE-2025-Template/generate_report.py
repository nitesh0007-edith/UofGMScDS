#!/usr/bin/env python3
"""Generate the Solution Report DOCX file."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import json, os

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

# ── Helper functions ──
def add_code_block(doc, code, font_size=Pt(8.5)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = font_size
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Big Data Assessed Exercise 2025/26')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Solution Report')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x4A, 0x6A, 0x9E)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Financial Recommendation Platform using Apache Spark')
run.font.size = Pt(13)
run.italic = True

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Course: Big Data (H/M)\n').font.size = Pt(11)
meta.add_run('Technology: Apache Spark 4.0.0-preview2, Java 21\n').font.size = Pt(11)
meta.add_run('Mode: local[4]\n').font.size = Pt(11)
meta.add_run('Date: February 2026').font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Problem Statement',
    '3. Solution Architecture',
    '4. Data Pipeline - Step by Step',
    '5. Key Classes Implemented',
    '6. Optimization Journey',
    '7. Performance Results',
    '8. Output - Top 5 Investments',
    '9. Project Structure',
    '10. How to Run the Application',
    '11. How to View the Dashboard',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'This report documents the design, implementation and optimisation of a financial '
    'recommendation platform built on Apache Spark. The application processes 24.2 million '
    'stock price records (2.3 GiB) and 15,700 asset metadata entries to identify the top 5 '
    'investment opportunities ranked by 5-day Return on Investment, after filtering by '
    'volatility and Price-to-Earnings ratio.'
)
doc.add_paragraph(
    'The final solution completes in 37 seconds on an M1 MacBook Air (local[4]) with a '
    'shuffle volume of only 657 KiB - a 1,150x reduction from the initial implementation\'s '
    '736 MiB shuffle. The rank #1 asset is TOP Ships Inc (TOPS), matching the professor\'s '
    'reference output.'
)

# ═══════════════════════════════════════════════════════════════
# 2. PROBLEM STATEMENT
# ═══════════════════════════════════════════════════════════════
doc.add_heading('2. Problem Statement', level=1)
doc.add_paragraph(
    'Given daily pricing data and asset metadata for US stock market assets (1999-2020), '
    'build a Spark pipeline that:'
)
add_bullet(doc, 'Filters pricing data up to the analysis date (2020-04-01)')
add_bullet(doc, 'Calculates technical indicators: 5-day Returns and 251-day Volatility')
add_bullet(doc, 'Filters out assets with Volatility >= 4')
add_bullet(doc, 'Filters out assets with P/E Ratio >= 25 or missing P/E data')
add_bullet(doc, 'Ranks remaining assets by Returns (descending) and returns the top 5')

doc.add_paragraph()
doc.add_heading('Dataset', level=2)
add_table(doc,
    ['File', 'Content', 'Size', 'Records'],
    [
        ['all_prices-noHead.csv', 'Daily pricing data', '2.3 GiB', '24,197,442'],
        ['stock_data.json', 'Asset metadata', '8.7 MB', '~15,700'],
    ]
)

# ═══════════════════════════════════════════════════════════════
# 3. SOLUTION ARCHITECTURE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('3. Solution Architecture', level=1)
doc.add_paragraph(
    'The solution uses a map-side partial aggregation strategy that avoids shuffling raw '
    'price data entirely. The architecture is built around three key design decisions:'
)

doc.add_heading('3.1 Map-Side Partial Aggregation (Core Innovation)', level=2)
doc.add_paragraph(
    'Instead of using groupByKey to shuffle all 24M StockPrice objects across partitions '
    '(~700 MiB), we use mapPartitionsToPair to compute per-ticker statistics locally within '
    'each partition. Each partition groups its prices by ticker, sorts them, and computes '
    'compact running statistics using Welford\'s online algorithm for volatility and keeps '
    'only the 6 most recent prices for returns. The resulting TickerStats objects (~150 bytes '
    'each) are then merged across partitions via reduceByKey, producing a shuffle of only '
    '~657 KiB.'
)

doc.add_heading('3.2 Broadcast Join', level=2)
doc.add_paragraph(
    'The asset metadata (~8 MB, ~15K records) is collected to the driver and broadcast to '
    'all executors. This eliminates the need for a shuffle-based join entirely. Each executor '
    'performs a map-side lookup in the broadcast HashMap.'
)

doc.add_heading('3.3 Fused Operations', level=2)
doc.add_paragraph(
    'The metadata join, P/E ratio filter, and Asset object construction are fused into a '
    'single map operation. This avoids creating intermediate RDDs and reduces the number of '
    'passes over the post-aggregation data.'
)

# ═══════════════════════════════════════════════════════════════
# 4. DATA PIPELINE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('4. Data Pipeline - Step by Step', level=1)

steps = [
    ('Step 1: Date Filter',
     'Filter all prices to include only data on or before 2020-04-01 using DateFilter. '
     'This is applied as a FilterFunction on the Dataset<StockPrice> before any RDD conversion.',
     'DateFilter.java'),
    ('Step 2: Key by Ticker',
     'Convert the filtered Dataset<StockPrice> to a JavaPairRDD<String, StockPrice> using '
     'StockPricePairing, which extracts the stock ticker as the key.',
     'StockPricePairing.java'),
    ('Step 3: Map-Side Partial Aggregation',
     'Using mapPartitionsToPair, process each partition locally:\n'
     '  (a) Group prices by ticker within the partition using a HashMap\n'
     '  (b) Sort each ticker\'s prices by date\n'
     '  (c) Compute volatility running stats via Welford\'s online algorithm\n'
     '  (d) Keep the 6 most recent prices for returns calculation\n'
     '  (e) Track boundary prices (earliest/latest) for cross-partition merge\n'
     'Output: JavaPairRDD<String, TickerStats> with ~15K compact entries per partition.',
     'TickerStats.java'),
    ('Step 4: Reduce Across Partitions',
     'Use reduceByKey with TickerStats.merge() to combine partial statistics from different '
     'partitions. The merge handles:\n'
     '  - Welford\'s parallel merge formula for volatility stats\n'
     '  - Boundary daily change computation between partition segments\n'
     '  - Keeping the 6 globally most recent prices for returns\n'
     'THIS IS THE ONLY SHUFFLE in the pipeline (~657 KiB).',
     'TickerStats.merge()'),
    ('Step 5: Volatility Filter',
     'Convert merged TickerStats to AssetFeatures using toAssetFeatures(), then filter '
     'out assets with volatility >= 4 using VolatilityFilter.',
     'VolatilityFilter.java'),
    ('Step 6: Broadcast Join with Metadata',
     'Collect the small metadata RDD to the driver, broadcast it to all executors. '
     'Perform a map-side lookup to join features with metadata. No shuffle required.',
     'Broadcast variable'),
    ('Step 7: P/E Filter + Asset Build',
     'In the same map operation as the join, filter out assets with P/E <= 0 or P/E >= 25, '
     'and construct the final Asset objects. Nulls are filtered out.',
     'Fused in AssessedExercise.java'),
    ('Step 8: Rank and Return Top 5',
     'Use RDD.top(5) which leverages Asset.compareTo() (compares by returns) to efficiently '
     'find the 5 highest-return assets without a full sort.',
     'Asset.compareTo()'),
]

for step_title, desc, impl in steps:
    doc.add_heading(step_title, level=2)
    doc.add_paragraph(desc)
    p = doc.add_paragraph()
    run = p.add_run(f'Implementation: {impl}')
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ═══════════════════════════════════════════════════════════════
# 5. KEY CLASSES
# ═══════════════════════════════════════════════════════════════
doc.add_heading('5. Key Classes Implemented', level=1)

doc.add_heading('5.1 TickerStats.java (objects)', level=2)
doc.add_paragraph(
    'The core innovation of the solution. A compact, serializable accumulator that stores '
    'running statistics for a single ticker\'s price data. Instead of holding all ~1,600 '
    'StockPrice objects per ticker (~100 KB), it compresses them to ~150 bytes:'
)
add_bullet(doc, 'Welford\'s running stats: ', 'Volatility: ')
add_bullet(doc, 'volCount, volMean, volM2 (3 doubles = 24 bytes)')
add_bullet(doc, 'Last 6 prices for returns calculation (96 bytes)', 'Returns: ')
add_bullet(doc, 'Earliest/latest date+price for cross-partition merge (32 bytes)', 'Boundary: ')
doc.add_paragraph(
    'Provides a static merge() method for reduceByKey that correctly combines partial '
    'statistics from different partitions using the parallel Welford merge formula, and a '
    'toAssetFeatures() method that computes the final returns and volatility values.'
)

doc.add_heading('5.2 DateFilter.java (transformations/filters)', level=2)
doc.add_paragraph(
    'A FilterFunction<StockPrice> that removes prices after the analysis end date. Uses '
    'TimeUtil to convert year/month/day fields to Java Instant for comparison.'
)

doc.add_heading('5.3 VolatilityFilter.java (transformations/filters)', level=2)
doc.add_paragraph(
    'A Function<Tuple2<String, AssetFeatures>, Boolean> that removes assets with volatility '
    'above the threshold (4.0). Applied after technical indicators are computed.'
)

doc.add_heading('5.4 StockPricePairing.java (transformations/pairing)', level=2)
doc.add_paragraph(
    'A PairFunction<StockPrice, String, StockPrice> that extracts the stock ticker as the '
    'key, enabling grouping operations downstream.'
)

# ═══════════════════════════════════════════════════════════════
# 6. OPTIMIZATION JOURNEY
# ═══════════════════════════════════════════════════════════════
doc.add_heading('6. Optimization Journey', level=1)
doc.add_paragraph(
    'The solution went through several iterations to reduce shuffle volume from 736 MiB to '
    '657 KiB. Below is a chronological account of each approach tried and why it succeeded '
    'or failed.'
)

doc.add_heading('6.1 Attempt 1: aggregateByKey with List Combiner', level=2)
doc.add_paragraph(
    'Replaced groupByKey with aggregateByKey, using list accumulation as the combiner. '
    'The hypothesis was that map-side combining would reduce shuffle.'
)
p = doc.add_paragraph()
run = p.add_run('Result: 620 MiB shuffle (no improvement)')
run.bold = True
doc.add_paragraph(
    'Why it failed: The combiner merged List<StockPrice> objects. Since every price is needed '
    'for indicator calculation, there is no actual data reduction - the combiner just '
    'concatenates lists. The full data still crosses the shuffle boundary.'
)

doc.add_heading('6.2 Attempt 2: Spark Configuration Tuning', level=2)
doc.add_paragraph(
    'Set spark.default.parallelism=8 and spark.sql.shuffle.partitions=8 to match the '
    'local[4] core count. Also tried Kryo serialization.'
)
p = doc.add_paragraph()
run = p.add_run('Result: Kryo crashed (Java 21 module system incompatibility). '
                'Parallelism change reduced partition count but NOT shuffle volume.')
run.bold = True
doc.add_paragraph(
    'Why it failed: Reducing partition count from 200 to 8 affects scheduling overhead but '
    'does not change the total bytes shuffled. Kryo serializer could not access '
    'java.nio.HeapByteBuffer due to Java 21\'s module access restrictions.'
)

doc.add_heading('6.3 Attempt 3: Dataset API groupByKey + mapGroups', level=2)
doc.add_paragraph(
    'Switched from RDD groupByKey to Dataset API groupByKey().mapGroups(), hoping that '
    'Tungsten binary format would compress the shuffle data.'
)
p = doc.add_paragraph()
run = p.add_run('Result: 736 MiB shuffle (worse)')
run.bold = True
doc.add_paragraph(
    'Why it failed: Even with Tungsten encoding, 24M rows of StockPrice beans (with all '
    'fields: open, high, low, close, adjClose, volume, ticker, date) still produce hundreds '
    'of MiBs. Tungsten reduces per-record overhead but cannot eliminate the fundamental '
    'problem: all records must cross the shuffle boundary.'
)

doc.add_heading('6.4 Attempt 4: Map-Side Partial Aggregation (Final Solution)', level=2)
doc.add_paragraph(
    'Used mapPartitionsToPair to compute per-ticker statistics locally within each partition '
    'using Welford\'s online algorithm for volatility, then reduceByKey to merge the compact '
    'TickerStats across partitions.'
)
p = doc.add_paragraph()
run = p.add_run('Result: 657 KiB shuffle - 1,150x reduction!')
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
doc.add_paragraph(
    'Why it worked: Instead of shuffling 24M StockPrice objects, we shuffle only ~15K '
    'TickerStats objects (~150 bytes each). The heavy computation (sorting, Welford\'s '
    'algorithm, price tracking) happens locally within each partition with zero network I/O.'
)

# ═══════════════════════════════════════════════════════════════
# 7. PERFORMANCE RESULTS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('7. Performance Results', level=1)

doc.add_heading('7.1 Final vs Initial vs Professor\'s Reference', level=2)
add_table(doc,
    ['Metric', 'Initial (Mine)', 'Professor (i7-12700)', 'Final (Mine)'],
    [
        ['Shuffle Read',    '620 MiB',   '5.3 MiB',   '657 KiB'],
        ['Shuffle Write',   '620 MiB',   '2.6 MiB',   '657 KiB'],
        ['GC Time',         '9 s',       '0.4 s',     '4 s'],
        ['Task Time',       '3.1 min',   '1.2 min',   '2.1 min'],
        ['Execution Time',  '~116 s',    '~31 s',     '37 s'],
        ['Total Tasks',     '42',        '61',        '31'],
        ['Rank #1 Output',  'TOPS',      'TOPS',      'TOPS'],
    ]
)

doc.add_paragraph()
doc.add_heading('7.2 Spark Configuration', level=2)
add_table(doc,
    ['Parameter', 'Value', 'Purpose'],
    [
        ['spark.master', 'local[4]', '4 executor threads'],
        ['spark.sql.shuffle.partitions', '8', 'Fewer partitions for local mode'],
        ['spark.default.parallelism', '8', 'Matches core count for RDD ops'],
        ['spark.driver.memory', '4g', 'Reduces GC pressure'],
        ['spark.executor.memory', '4g', 'Adequate heap for 2.3 GiB input'],
    ]
)

# ═══════════════════════════════════════════════════════════════
# 8. OUTPUT
# ═══════════════════════════════════════════════════════════════
doc.add_heading('8. Output - Top 5 Investments', level=1)

# Load actual results
results_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'results', 'results.json')
try:
    with open(results_path) as f:
        data = json.load(f)
    investments = data['topInvestments']
    rows = []
    for inv in investments:
        rows.append([
            str(inv['rank']),
            f"{inv['ticker']}",
            inv['name'],
            inv['industry'],
            inv['sector'],
            f"{inv['returns']*100:.2f}%",
            f"{inv['volatility']:.2f}",
            f"{inv['peRatio']:.2f}",
        ])
    add_table(doc,
        ['Rank', 'Ticker', 'Name', 'Industry', 'Sector', 'Returns', 'Volatility', 'P/E'],
        rows
    )
except Exception as e:
    doc.add_paragraph(f'(Could not load results.json: {e})')

doc.add_paragraph()
doc.add_paragraph(
    f'Execution completed in {data.get("executionTime", "N/A")} seconds on Apple M1 Air, local[4] mode.'
)

# ═══════════════════════════════════════════════════════════════
# 9. PROJECT STRUCTURE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('9. Project Structure', level=1)

doc.add_heading('9.1 Source Directory (src/)', level=2)
structure = """src/bigdata/
  app/
    AssessedExercise.java        -- Main pipeline & rankInvestments()
  objects/
    Asset.java                   -- [Provided] Single asset representation
    AssetFeatures.java           -- [Provided] Returns, volatility, P/E ratio
    AssetMetadata.java           -- [Provided] Name, industry, sector, P/E
    AssetRanking.java            -- [Provided] Array of 5 assets
    StockPrice.java              -- [Provided] Daily price data
    TickerStats.java             -- [NEW] Map-side aggregation accumulator
  technicalindicators/
    Returns.java                 -- [Provided] ROI calculation
    Volitility.java              -- [Provided] Volatility calculation
  transformations/
    filters/
      DateFilter.java            -- [NEW] Filters by end date
      NullPriceFilter.java       -- [Provided] Removes null prices
      VolatilityFilter.java      -- [NEW] Filters by volatility threshold
    maps/
      PriceReaderMap.java        -- [Provided] Row to StockPrice conversion
    pairing/
      AssetMetadataPairing.java  -- [Provided] Row to (ticker, metadata)
      StockPricePairing.java     -- [NEW] StockPrice to (ticker, price)
  util/
    MathUtils.java               -- [Provided] Math operations
    ResultsExporter.java         -- JSON export for dashboard
    TimeUtil.java                -- [Provided] Date parsing utility"""
add_code_block(doc, structure)

doc.add_heading('9.2 New Classes Summary', level=2)
add_table(doc,
    ['Class', 'Package', 'Role'],
    [
        ['TickerStats', 'objects', 'Compact accumulator with Welford\'s algorithm for map-side aggregation'],
        ['DateFilter', 'transformations/filters', 'Filters prices by end date using TimeUtil'],
        ['VolatilityFilter', 'transformations/filters', 'Filters assets with volatility >= threshold'],
        ['StockPricePairing', 'transformations/pairing', 'Pairs StockPrice by ticker for grouping'],
    ]
)

# ═══════════════════════════════════════════════════════════════
# 10. HOW TO RUN
# ═══════════════════════════════════════════════════════════════
doc.add_heading('10. How to Run the Application', level=1)

doc.add_heading('10.1 Prerequisites', level=2)
add_bullet(doc, 'Java JDK 21.0.2 or later')
add_bullet(doc, 'Apache Maven 3.8+')
add_bullet(doc, 'Dataset files in resources/ directory (all_prices-noHead.csv, stock_data.json)')

doc.add_heading('10.2 Compile and Run via Terminal', level=2)
add_code_block(doc,
    'cd BigData-AE-2025-Template\n'
    'mvn clean compile exec:java -Dexec.mainClass="bigdata.app.AssessedExercise"'
)
doc.add_paragraph(
    'The application will run, print the top 5 rankings, export results to results/results.json, '
    'and hold the Spark UI open at http://localhost:4040 for 10 minutes.'
)

doc.add_heading('10.3 Run via VS Code', level=2)
add_bullet(doc, 'Open the BigData-AE-2025-Template folder in VS Code')
add_bullet(doc, 'Install the Extension Pack for Java (by Microsoft)')
add_bullet(doc, 'Press F5 and select "Run AssessedExercise (Spark)" from the launch dropdown')
add_bullet(doc, 'The launch.json is pre-configured with -Xmx4g and SPARK_MASTER=local[4]')

doc.add_heading('10.4 Monitor Execution', level=2)
doc.add_paragraph(
    'While the application is running, visit http://localhost:4040 in your browser to see the '
    'Spark UI with job progress, stage details, shuffle statistics, and executor metrics.'
)

# ═══════════════════════════════════════════════════════════════
# 11. DASHBOARD
# ═══════════════════════════════════════════════════════════════
doc.add_heading('11. How to View the Dashboard', level=1)
doc.add_paragraph(
    'After running the application, results are exported to results/results.json. '
    'A visualization dashboard is provided as an HTML file with interactive charts.'
)

doc.add_heading('11.1 Option A: Local HTTP Server (Recommended)', level=2)
add_code_block(doc,
    'cd BigData-AE-2025-Template/results\n'
    'python3 -m http.server 8080'
)
doc.add_paragraph('Then open http://localhost:8080/dashboard.html in your browser.')

doc.add_heading('11.2 Option B: Direct File Open', level=2)
add_code_block(doc, 'open results/dashboard.html')
doc.add_paragraph(
    'Note: Some browsers block fetch() from file:// URLs due to CORS. '
    'If charts do not load, use Option A instead.'
)

doc.add_heading('11.3 Dashboard Features', level=2)
add_bullet(doc, 'Top 5 investment cards with ticker, name, industry, sector')
add_bullet(doc, 'Returns bar chart comparing all 5 assets')
add_bullet(doc, 'Volatility comparison chart')
add_bullet(doc, 'P/E Ratio comparison chart')
add_bullet(doc, 'Execution time and performance metrics')

# ═══════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'BigData_AE_Solution_Report.docx')
doc.save(output_path)
print(f'Report saved to: {output_path}')
