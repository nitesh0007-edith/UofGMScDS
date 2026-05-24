"""Build the final DOCX report for our IV group project.

Reads evaluation data from the xlsx, generates analysis charts, embeds
dashboard screenshots, and writes everything into a formatted Word doc.
"""

import os
import io
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
EVAL_PATH = os.path.join(BASE_DIR, "evaluation", "Evaluation_Data_Collection.xlsx")
REPORT_DIR = os.path.join(BASE_DIR, "report")
OUTPUT_DOCX = os.path.join(BASE_DIR, "output", "IV_Group_Project_Report.docx")

SCREENSHOT_A = os.path.join(REPORT_DIR, "System_A_screenshot.png")
SCREENSHOT_B = os.path.join(REPORT_DIR, "System_B_screenshot.png")
SCREENSHOT_C = os.path.join(REPORT_DIR, "System_C_screenshot.png")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex):
    """Apply a background fill colour to a Word table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shd)


def add_heading(doc, text, level=1):
    """Add a styled heading in the UofG navy colour."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x00, 0x38, 0x65)


def add_body(doc, text, bold=False):
    """Add a normal paragraph of text."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15


def add_bullet(doc, text):
    """Add a bullet-point item."""
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)


def add_caption(doc, text):
    """Centred italic caption for figures."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.paragraph_format.space_after = Pt(10)


def add_screenshot(doc, path, caption, width=6.0):
    """Insert a screenshot image followed by a caption."""
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
        add_caption(doc, caption)
    else:
        add_body(doc, f"[Screenshot not found: {os.path.basename(path)}]")


def save_chart_to_buffer(fig):
    """Render a matplotlib figure to an in-memory PNG."""
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    plt.close(fig)
    return buf


def add_chart_image(doc, buf, caption, width=5.5):
    """Insert a chart image from a buffer into the doc."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(width))
    add_caption(doc, caption)


def make_table(doc, headers, rows):
    """Insert a nicely formatted table with headers and data rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, "003865")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()  # spacing


# ---------------------------------------------------------------------------
# Load evaluation data
# ---------------------------------------------------------------------------

def load_eval_data():
    """Pull in the evaluation sheets from the xlsx."""
    df = pd.read_excel(EVAL_PATH, sheet_name='Data Collection')
    rank_df = pd.read_excel(EVAL_PATH, sheet_name='Post-Study Ranking')
    task_ref = pd.read_excel(EVAL_PATH, sheet_name='Task Reference')
    return df, rank_df, task_ref


# ---------------------------------------------------------------------------
# Chart generators for the evaluation section
# ---------------------------------------------------------------------------

def create_time_by_system_task(df):
    """Grouped bar chart comparing average completion times."""
    pivot = df.pivot_table(values='Time (seconds)', index='Task',
                           columns='System', aggfunc='mean')
    pivot = pivot[['A', 'B', 'C']]

    fig, ax = plt.subplots(figsize=(7, 4))
    pivot.plot(kind='bar', ax=ax, color=['#e45756', '#4c78a8', '#54a24b'],
               edgecolor='white', width=0.7)
    ax.set_ylabel('Avg Completion Time (seconds)', fontsize=10)
    ax.set_xlabel('Task', fontsize=10)
    ax.set_title('Average Task Completion Time by System', fontsize=12, fontweight='bold')
    ax.legend(title='System', frameon=True, fontsize=9)
    ax.set_xticklabels(pivot.index, rotation=0)
    ax.grid(axis='y', alpha=0.3)
    ax.set_axisbelow(True)
    plt.tight_layout()
    return save_chart_to_buffer(fig)


def create_accuracy_by_system_task(df):
    """Grouped bar chart comparing average accuracy scores."""
    pivot = df.pivot_table(values='Accuracy (0/1/2)', index='Task',
                           columns='System', aggfunc='mean')
    pivot = pivot[['A', 'B', 'C']]

    fig, ax = plt.subplots(figsize=(7, 4))
    pivot.plot(kind='bar', ax=ax, color=['#e45756', '#4c78a8', '#54a24b'],
               edgecolor='white', width=0.7)
    ax.set_ylabel('Avg Accuracy (0-2)', fontsize=10)
    ax.set_xlabel('Task', fontsize=10)
    ax.set_title('Average Accuracy by System and Task', fontsize=12, fontweight='bold')
    ax.legend(title='System', frameon=True, fontsize=9)
    ax.set_xticklabels(pivot.index, rotation=0)
    ax.set_ylim(0, 2.2)
    ax.grid(axis='y', alpha=0.3)
    ax.set_axisbelow(True)
    plt.tight_layout()
    return save_chart_to_buffer(fig)


def create_likert_by_system(df):
    """Simple bar chart of average ease-of-use ratings."""
    means = df.groupby('System')['Likert (1-5)'].mean().reindex(['A', 'B', 'C'])

    fig, ax = plt.subplots(figsize=(5, 3.5))
    bars = ax.bar(means.index, means.values,
                  color=['#e45756', '#4c78a8', '#54a24b'], edgecolor='white', width=0.5)
    ax.set_ylabel('Avg Likert Rating (1-5)', fontsize=10)
    ax.set_xlabel('System', fontsize=10)
    ax.set_title('Average Ease-of-Use Rating by System', fontsize=12, fontweight='bold')
    ax.set_ylim(0, 5.5)
    ax.grid(axis='y', alpha=0.3)
    ax.set_axisbelow(True)
    for bar, val in zip(bars, means.values):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.1,
                f'{val:.2f}', ha='center', fontsize=10, fontweight='bold')
    plt.tight_layout()
    return save_chart_to_buffer(fig)


def create_ranking_chart(rank_df):
    """Show how participants ranked each system (best/middle/worst)."""
    ranking_data = {'A': [0, 0, 0], 'B': [0, 0, 0], 'C': [0, 0, 0]}
    for _, row in rank_df.iterrows():
        r1 = row['Rank 1 (Best System)'].replace('System ', '')
        r2 = row['Rank 2'].replace('System ', '').strip()
        r3 = row['Rank 3 (Worst System)'].replace('System ', '').strip()
        ranking_data[r1][0] += 1
        ranking_data[r2][1] += 1
        ranking_data[r3][2] += 1

    fig, ax = plt.subplots(figsize=(5, 3.5))
    x = np.arange(3)
    width = 0.25
    colors_rank = ['#2ca02c', '#f58518', '#d62728']

    for i, (rank_label, color) in enumerate(zip(
        ['Rank 1 (Best)', 'Rank 2', 'Rank 3 (Worst)'], colors_rank
    )):
        vals = [ranking_data[s][i] for s in ['A', 'B', 'C']]
        ax.bar(x + i * width, vals, width, label=rank_label, color=color, edgecolor='white')

    ax.set_xticks(x + width)
    ax.set_xticklabels(['System A', 'System B', 'System C'])
    ax.set_ylabel('Number of Participants', fontsize=10)
    ax.set_title('Post-Study System Rankings', fontsize=12, fontweight='bold')
    ax.legend(fontsize=8, frameon=True)
    ax.yaxis.set_major_locator(ticker.MaxNLocator(integer=True))
    ax.grid(axis='y', alpha=0.3)
    ax.set_axisbelow(True)
    plt.tight_layout()
    return save_chart_to_buffer(fig)


def create_accuracy_heatmap(df):
    """Colour-coded grid of accuracy scores by system and task."""
    pivot = df.pivot_table(values='Accuracy (0/1/2)', index='System',
                           columns='Task', aggfunc='mean')
    pivot = pivot.reindex(['A', 'B', 'C'])

    fig, ax = plt.subplots(figsize=(6, 3))
    im = ax.imshow(pivot.values, cmap='RdYlGn', aspect='auto', vmin=0, vmax=2)
    ax.set_xticks(range(len(pivot.columns)))
    ax.set_xticklabels(pivot.columns)
    ax.set_yticks(range(len(pivot.index)))
    ax.set_yticklabels([f'System {s}' for s in pivot.index])
    ax.set_title('Accuracy Heatmap (System x Task)', fontsize=12, fontweight='bold')

    for i in range(len(pivot.index)):
        for j in range(len(pivot.columns)):
            val = pivot.values[i, j]
            ax.text(j, i, f'{val:.2f}', ha='center', va='center',
                    fontsize=10, fontweight='bold',
                    color='white' if val < 1.0 else 'black')

    plt.colorbar(im, ax=ax, label='Avg Accuracy')
    plt.tight_layout()
    return save_chart_to_buffer(fig)


# ---------------------------------------------------------------------------
# Report builder
# ---------------------------------------------------------------------------

def build_report():
    """Main function — assembles the entire report document."""
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Load eval data
    eval_df, rank_df, task_ref = load_eval_data()

    # =====================================================================
    # TITLE PAGE
    # =====================================================================

    for _ in range(4):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('IV Group Project:\nMultiview Visualisation')
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x38, 0x65)

    doc.add_paragraph()

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run('Information Visualisation — University of Glasgow')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run('March 2026')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    # YouTube link at the TOP
    yt_p = doc.add_paragraph()
    yt_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = yt_p.add_run('Demo Video: ')
    run.font.size = Pt(11)
    run.bold = True
    run = yt_p.add_run('[INSERT YOUTUBE LINK HERE]')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x00, 0x55, 0xAA)
    run.underline = True

    doc.add_page_break()

    # =====================================================================
    # PART 1: THE DATA (max 400 words)
    # =====================================================================

    add_heading(doc, '1. The Data', level=1)

    add_body(doc,
        'Title: Glasgow Weather Dataset (2015-2019) — A five-year daily record of '
        'weather conditions in Glasgow, United Kingdom.'
    , bold=True)

    add_body(doc,
        'Source: The dataset was obtained from Kaggle '
        '(https://www.kaggle.com/datasets/phyamal/glasgow-weather-data-20152019) and contains '
        '1,795 daily observations spanning from 1 January 2015 to 30 November 2019. Each row '
        'represents a single day, making this a tabular, flat-file dataset with no hierarchical '
        'nesting in its raw form.'
    )

    add_body(doc,
        'Data categorisation (using Munzner, 2014, Chapter 2 terminology):'
    , bold=True)

    add_bullet(doc,
        'Dataset type: Table (flat table with items and attributes). Each item is a day; '
        'each attribute is a measured weather variable.'
    )
    add_bullet(doc,
        'Attribute types: The dataset contains quantitative attributes (tempMin, tempMax, '
        'windSpeed, humidity, visibility, cloudCover — all continuous, ratio-scale measurements), '
        'categorical/nominal attributes (desc — six weather categories: rain, clear-day, '
        'partly-cloudy-day, cloudy, fog, unknown; and summary — free-text descriptions), '
        'and a temporal/ordered attribute (day — sequential dates).'
    )
    add_bullet(doc,
        'Derived attributes: We computed several additional attributes to support our tasks: '
        'tempRange (quantitative, derived as tempMax minus tempMin), avgTemp (quantitative), '
        'month and year (ordinal, extracted from day), season (categorical, mapped from month), '
        'monthName (ordinal), and dayOfYear (quantitative sequential).'
    )
    add_bullet(doc,
        'Ordering direction: Temperature, wind speed, and visibility are diverging in nature '
        '(both high and low values carry meaning), while humidity and cloud cover are sequential '
        '(0 to 1 ratio scale). The temporal attribute day provides a natural sequential ordering.'
    )
    add_bullet(doc,
        'Data quality: The raw dataset contained 7 null values in summary (filled with '
        '"No description"), 184 blank values in desc (filled with "unknown"), and 3 null values '
        'in cloudCover (filled with the column median of 0.71). No outliers were removed.'
    )

    add_body(doc,
        'The dataset is well suited for multiview visualisation because it combines temporal, '
        'quantitative, and categorical dimensions, enabling exploration of seasonal trends, '
        'weather-type distributions, and correlations between atmospheric variables. Its '
        'moderate size (1,795 items, 9 original attributes plus 7 derived) is manageable for '
        'interactive client-side rendering while being rich enough to reveal meaningful patterns.'
    )

    # =====================================================================
    # PART 2: THE TASKS (max 400 words)
    # =====================================================================

    add_heading(doc, '2. The Tasks', level=1)

    add_body(doc,
        'We defined five tasks that cover a range of the task taxonomy from Munzner (2014, '
        'Chapter 3). Each task requires users to perform different actions on different targets, '
        'ensuring variety in the cognitive operations demanded by our systems.'
    )

    tasks = [
        ('T1', 'Identify the season or month with the highest average temperature range.',
         'Action: Identify (find an extreme value). Target: One attribute (tempRange) across '
         'a categorical grouping (season/month). This is a lookup task requiring the user to '
         'compare aggregated values and locate the maximum. It tests the system\'s ability to '
         'support comparison of means across groups.'),

        ('T2', 'Find the top 3 weather types (desc) by frequency across the dataset.',
         'Action: Identify + Rank (find and order extremes). Target: Distribution of a '
         'categorical attribute (desc). This requires counting occurrences and sorting, '
         'testing whether the system supports frequency-based ranking of categories effectively.'),

        ('T3', 'Determine which year had the highest average wind speed.',
         'Action: Identify (find an extreme). Target: One quantitative attribute (windSpeed) '
         'across a temporal grouping (year). Similar to T1 but uses a different attribute and '
         'temporal rather than seasonal grouping.'),

        ('T4', 'Find the weather type with the highest average visibility, filtered by a '
         'minimum humidity threshold (humidity > 0.85).',
         'Action: Filter + Identify. Target: One attribute (visibility) within a filtered '
         'subset (humidity > 0.85). This compound task tests the system\'s ability to support '
         'interactive data selection followed by identification of an extreme, which is critical '
         'for evaluating slider-based filtering.'),

        ('T5', 'Identify months where wind speed is above a user-set threshold AND humidity '
         'is above a user-set threshold. Find the top 3 seasons with the highest average '
         'cloud cover.',
         'Action: Filter + Identify + Rank. Target: Multiple attributes with compound filtering. '
         'This is the most complex task, requiring simultaneous adjustment of two slider thresholds '
         'followed by identifying and ranking seasonal aggregates. It tests the system\'s ability '
         'to support multi-constraint exploration.'),
    ]

    for tid, desc, analysis in tasks:
        p = doc.add_paragraph()
        run = p.add_run(f'{tid}: ')
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(desc)
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)

        p2 = doc.add_paragraph()
        run = p2.add_run(analysis)
        run.font.size = Pt(10)
        run.italic = True
        p2.paragraph_format.space_after = Pt(8)
        p2.paragraph_format.left_indent = Pt(18)

    add_body(doc,
        'These tasks collectively cover the actions of identify, compare, rank, and filter, '
        'applied to targets including individual attributes, distributions, extremes, and '
        'filtered subsets (Shneiderman, 1996). Task T4 specifically involves selection of a '
        'data subset, which is a prerequisite for Part 4 (Generalised Selection).'
    )

    # =====================================================================
    # PART 3: THE SYSTEMS (code submission — describe here)
    # =====================================================================

    add_heading(doc, '3. The Systems', level=1)

    add_body(doc,
        'Three visualisation systems were implemented using Python 3.10 and Altair 5.x. '
        'Each system is a self-contained interactive HTML dashboard with multiple linked '
        'Vega-Lite charts. All three systems support all five tasks (T1-T5) and share a '
        'common data preparation module (data_prep.py). The code is submitted as three '
        'clearly labelled zipped folders (A, B, C).'
    )

    # System A
    add_heading(doc, 'System A: Category + Attribute Overview', level=2)
    add_body(doc,
        'System A provides an overview of weather patterns across categories. It uses a warm '
        'colour palette (oranges, salmons) and features four charts in a 2x2 grid:'
    )
    add_bullet(doc,
        'Top-left: Horizontal bar chart ranking weather types by a user-selected metric '
        '(dropdown with 7 options including Count).')
    add_bullet(doc,
        'Top-right: Dynamic bar chart showing distribution by month, season, or year '
        '(controlled by a Group By dropdown), filtered by bar click from chart 1.')
    add_bullet(doc,
        'Bottom-left: Scatter plot of Date vs Wind Speed with bidirectional brush linking '
        'to chart 4.')
    add_bullet(doc,
        'Bottom-right: Scatter plot of Humidity vs Visibility with bidirectional brush '
        'linking to chart 3.')
    add_bullet(doc,
        'Controls: 2 dropdowns (Rank by, Group by) and 3 sliders (min humidity, min wind '
        'speed, max cloud cover).')
    add_bullet(doc,
        'Interactions: Click selection on bars, bidirectional interval brushing between '
        'scatter plots, scroll-to-zoom on scatter plots.')

    add_screenshot(doc, SCREENSHOT_A,
        'Figure 1: System A — Category + Attribute Overview dashboard')

    # System B
    add_heading(doc, 'System B: Temporal Decomposition + Heatmap', level=2)
    add_body(doc,
        'System B focuses on temporal patterns, using a blue/green/yellow colour treatment '
        'and featuring four charts:'
    )
    add_bullet(doc,
        'Top-left: Horizontal bar chart with dynamic grouping (month/season/weather type/year) '
        'and metric selection.')
    add_bullet(doc,
        'Top-right: Calendar heatmap (Month x Year) displaying the mean of the selected metric '
        'with a viridis colour scheme.')
    add_bullet(doc,
        'Bottom-left: Grouped bar chart comparing wind speed and humidity side by side for '
        'each season.')
    add_bullet(doc,
        'Bottom-right: Generalised selection chart — line chart showing yearly weather type '
        'counts, with drill-down capability (see Part 4).')
    add_bullet(doc,
        'Controls: 3 dropdowns (Metric, Group by, Generalisation Level) and 2 sliders '
        '(wind speed, humidity).')
    add_bullet(doc,
        'Interactions: Click filtering across charts, bidirectional season click, heatmap '
        'cell click filtering, generalised selection drill-down.')

    add_screenshot(doc, SCREENSHOT_B,
        'Figure 2: System B — Temporal Decomposition + Heatmap dashboard')

    # System C
    add_heading(doc, 'System C: Full Exploration with Legend Filtering', level=2)
    add_body(doc,
        'System C provides maximum exploration flexibility with a multi-hue palette, '
        'using both colour (weather type) and shape (season) channels:'
    )
    add_bullet(doc,
        'Left: Horizontal bar chart ranking groups by metric with vivid per-group colours.')
    add_bullet(doc,
        'Centre: Large scatter plot (Date vs Selected Metric) using colour for weather type '
        'and shape for season (circle=Winter, diamond=Spring, square=Summer, triangle=Autumn).')
    add_bullet(doc,
        'Bottom-left: Heatmap of monthly weather composition (Month x Year) linked to '
        'scatter brush.')
    add_bullet(doc,
        'Bottom-right: Season summary bar chart showing average cloud cover per season.')
    add_bullet(doc,
        'Controls: 2 dropdowns (Metric, Group bars by) and 3 sliders (min humidity, min wind '
        'speed, max cloud cover).')
    add_bullet(doc,
        'Interactions: Clickable colour legend (weather type filtering), clickable shape legend '
        '(season filtering), interval brush on scatter, bidirectional heatmap click, '
        'scroll-to-zoom.')

    add_screenshot(doc, SCREENSHOT_C,
        'Figure 3: System C — Full Exploration with Legend Filtering dashboard')

    # =====================================================================
    # PART 4: GENERALISED SELECTION (max 400 words)
    # =====================================================================

    add_heading(doc, '4. Generalised Selection', level=1)

    add_body(doc, 'Semantic Structure', bold=True)
    add_body(doc,
        'We defined a temporal hierarchy over our dataset with five levels of abstraction:'
    )
    add_bullet(doc, 'Level 0 (most specific): Individual Day — e.g., 15 July 2017')
    add_bullet(doc, 'Level 1: Month — e.g., July 2017')
    add_bullet(doc, 'Level 2: Season — e.g., Summer 2017')
    add_bullet(doc, 'Level 3: Year — e.g., 2017')
    add_bullet(doc, 'Level 4 (most general): Entire Dataset — 2015-2019')

    add_body(doc,
        'This hierarchy is semantically meaningful because weather patterns naturally cluster '
        'at each level: daily variation, monthly means, seasonal cycles, and year-on-year trends. '
        'The structure follows a containment relationship (Munzner, 2014): each day belongs to '
        'exactly one month, each month to one season, and each season to one year.'
    )

    add_body(doc, 'Traversal Policy', bold=True)
    add_body(doc,
        'Our traversal policy moves from general to specific (top-down navigation). The user '
        'begins at Level 4 (entire dataset overview), selects a year to move to Level 3, and '
        'then uses a dropdown to choose whether to view the selected year decomposed by season '
        '(Level 2) or by month (Level 1). This is not merely changing a global filter — it is '
        'a true generalised selection where the user selects an item at one level and the system '
        'responds by showing the corresponding items at a more detailed level of the hierarchy.'
    )

    add_body(doc, 'Implementation', bold=True)
    add_body(doc,
        'Generalised selection is implemented in System B\'s bottom-right chart. The chart '
        'uses an Altair layered composition with two marks:'
    )
    add_bullet(doc,
        'Layer A (Overview): A line chart with points showing the count of days per weather '
        'type per year. This represents Level 3-4 of the hierarchy. When no year is selected, '
        'all lines are fully opaque, providing the complete overview.'
    )
    add_bullet(doc,
        'Layer B (Detail): A bar chart that only appears when a year is selected (via click on '
        'a point in the overview). It filters to the clicked year and groups the data by the '
        'user-selected generalisation level (year, season, or month), controlled by the '
        '"Generalisation Level" dropdown.'
    )
    add_body(doc,
        'When the user clicks a year on the overview, the line chart fades (opacity drops to '
        '0.15) and the detail bars appear, showing the breakdown for that year. The user can '
        'then switch the dropdown from "Yearly Overview" to "By Season" or "By Month" to drill '
        'deeper into the hierarchy. Double-clicking resets back to the full overview. This '
        'interaction directly implements generalised selection as described in the Data Selection '
        'lecture: a selected object at one level is expanded into its constituent objects at a '
        'lower level through a single interaction.'
    )

    # =====================================================================
    # PART 5: DEMO VIDEO
    # =====================================================================

    add_heading(doc, '5. Demo Video', level=1)
    add_body(doc,
        'A demonstration video (under 5 minutes) showcasing all three systems, their design '
        'rationale, key interactions, and how they support the five tasks is available at:'
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[INSERT YOUTUBE LINK HERE]')
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x55, 0xAA)
    run.underline = True
    p.paragraph_format.space_after = Pt(12)

    # =====================================================================
    # PART 6: DESIGN COMPARISON (max 1200 words, 200 per decision)
    # =====================================================================

    add_heading(doc, '6. Design Comparison', level=1)

    add_body(doc,
        'We compare six design decisions across the three systems, drawing on IV theory to '
        'justify each choice.'
    )

    # Decision 1
    add_heading(doc, 'Decision 1: Primary Mark Type for the Overview Chart', level=2)
    add_body(doc,
        'All three systems feature a horizontal bar chart as the primary ranking chart. '
        'In System A, we use coloured bars mapped to weather type via the shared WEATHER_COLORS '
        'scale — clicking a bar applies a point selection that filters all other charts. '
        'System B uses a monochromatic blue bar chart where the selected bar highlights in darker '
        'blue, keeping visual emphasis on the temporal heatmap alongside. System C assigns a '
        'distinct vivid colour per group (rainbow scheme) regardless of the grouping dimension, '
        'maximising pop-out for individual categories. According to Mackinlay\'s (1986) '
        'effectiveness principle, position along a common scale (bar length) is the most effective '
        'channel for quantitative comparison, which all three systems leverage. However, the colour '
        'treatment differs: System A\'s consistent weather-type palette provides identity continuity '
        'across charts, while System C\'s distinct hues maximise discriminability. Best choice: '
        'System A, because its consistent colour mapping across charts reduces cognitive load and '
        'supports the Gestalt principle of similarity — the same weather type is recognisable '
        'everywhere by colour.'
    )

    # Decision 2
    add_heading(doc, 'Decision 2: Supporting Task T1 (Highest Avg Temperature Range)', level=2)
    add_body(doc,
        'For T1, System A uses a "Group by" dropdown that switches the secondary bar chart '
        'between month, season, and year groupings, combined with a "Rank by" dropdown set to '
        'Avg Temp Range. This idiom (Munzner, 2014) separates selection from display, letting '
        'the user explicitly choose the aggregation level. System B uses the same dropdown '
        'mechanism but adds a heatmap that provides spatial context for temporal patterns — the '
        'heatmap can show tempRange as a colour-coded Month x Year grid, allowing users to spot '
        'seasonal trends visually. System C relies on the scatter plot to show individual data '
        'points with tempRange on the Y-axis, requiring the user to mentally aggregate by season '
        'via shape encoding. Best choice: System B, because the combination of sorted bar chart '
        'and heatmap provides both precise ranking (bar length) and overview context (colour '
        'saturation), following Shneiderman\'s (1996) mantra of "overview first, details on demand."'
    )

    # Decision 3
    add_heading(doc, 'Decision 3: Linking Strategy Between Charts', level=2)
    add_body(doc,
        'System A implements bidirectional brushing between its two scatter plots: dragging on '
        'the Date vs Wind Speed scatter highlights corresponding points in the Humidity vs '
        'Visibility scatter, and vice versa. This follows the coordinated multiple views paradigm, '
        'allowing users to discover correlations across attribute spaces. System B uses '
        'unidirectional click-based linking: clicking a bar filters the heatmap, clicking a '
        'heatmap cell filters the charts below, and clicking a season in the grouped bar filters '
        'charts 1 and 2 (bidirectional). System C uses legend-based filtering: clicking a legend '
        'item globally filters all charts by weather type or season. Best choice: System C, '
        'because legend-based filtering is the most intuitive (users naturally associate legend '
        'items with categories) and the filtering is truly global, affecting every chart '
        'simultaneously (Gestalt principle of connectedness).'
    )

    # Decision 4
    add_heading(doc, 'Decision 4: Colour Channel Usage', level=2)
    add_body(doc,
        'System A uses a fixed categorical colour scale (WEATHER_COLORS) consistently mapping '
        'weather types to the same hue across all charts, following the expressiveness principle '
        '(Mackinlay, 1986) — categorical data encoded with categorical colour. System B uses '
        'sequential colour (viridis scheme) for its heatmap to encode quantitative metric values, '
        'and distinct colours for the line chart. This mixed approach provides appropriate '
        'colour treatment per chart type. System C maps colour to weather type and shape to season, '
        'using two visual channels simultaneously. While this doubles the encoded dimensions, it '
        'risks visual clutter (participants noted confusion with shapes and colours). Best choice: '
        'System A, because a single consistent categorical palette reduces the integral-separable '
        'dimension conflict and keeps the colour channel focused on one variable.'
    )

    # Decision 5
    add_heading(doc, 'Decision 5: Filter Mechanism for Task T4', level=2)
    add_body(doc,
        'T4 requires filtering by humidity threshold, then identifying the top weather type '
        'by visibility. System A uses a slider (binding_range) positioned below the charts that '
        'reactively filters all four charts simultaneously. System B also uses a slider, but the '
        'user must additionally change the Group By dropdown to "By Weather Type" and the Metric '
        'to "Avg Visibility," requiring three interactions. System C uses a slider combined with '
        'legend filtering — the user can set the humidity threshold and then click on individual '
        'weather types in the legend to compare. Best choice: System A, because the slider '
        'directly filters all charts including the bar chart that already shows weather types, '
        'requiring only two interactions (adjust slider, read bar chart) compared to three or '
        'more in other systems.'
    )

    # Decision 6
    add_heading(doc, 'Decision 6: Temporal Exploration Idiom', level=2)
    add_body(doc,
        'System A represents time as a continuous axis in a scatter plot (Date vs Wind Speed), '
        'allowing fine-grained inspection of individual days but requiring users to mentally '
        'aggregate over time periods. System B uses a calendar heatmap (Month x Year rect marks) '
        'that spatialises time into a grid, enabling immediate recognition of seasonal and annual '
        'patterns through Gestalt proximity. The heatmap encodes the mean of the selected metric '
        'as colour intensity, supporting "overview first" exploration. System C also uses a scatter '
        'with time on X but adds shape encoding for season, requiring users to distinguish four '
        'shapes in a dense point cloud. Best choice: System B, because the heatmap provides an '
        'immediate spatial overview of temporal patterns that leverages pre-attentive processing '
        'of colour saturation, whereas scatter plots require effortful scanning of individual '
        'points.'
    )

    # =====================================================================
    # PART 7: USER EVALUATION (max 1000 words)
    # =====================================================================

    add_heading(doc, '7. User Evaluation Comparison', level=1)

    add_heading(doc, '7.1 Methodology', level=2)
    add_body(doc,
        'We conducted a within-subjects user evaluation with five participants (P1-P5). Each '
        'participant completed all five tasks (T1-T5) on all three systems (A, B, C), resulting '
        'in 75 total observations (5 participants x 3 systems x 5 tasks). System presentation '
        'order was counterbalanced using a Latin square design to mitigate learning effects:'
    )
    add_bullet(doc, 'P1, P4: A -> B -> C')
    add_bullet(doc, 'P2, P5: B -> C -> A')
    add_bullet(doc, 'P3: C -> A -> B')

    add_body(doc,
        'For each (participant, system, task) combination, we recorded three quantitative '
        'measures:'
    )
    add_bullet(doc,
        'Time (seconds): Task completion time, measured from when the participant began '
        'interacting to when they gave their answer.')
    add_bullet(doc,
        'Accuracy (0/1/2): 0 = wrong or could not complete; 1 = partially correct; '
        '2 = fully correct.')
    add_bullet(doc,
        'Likert (1-5): Subjective ease-of-use rating (1 = very difficult, 5 = very easy).')

    add_body(doc,
        'After completing all tasks on all systems, each participant provided a post-study '
        'ranking of the three systems from best (1) to worst (3), along with comments explaining '
        'their preference. The evaluator read each task aloud, did not guide participants toward '
        'answers, and recorded observations in free-text notes.'
    )

    add_heading(doc, '7.2 Results', level=2)

    # Summary table
    add_body(doc, 'Overall Performance Summary', bold=True)
    sys_summary = []
    for sys in ['A', 'B', 'C']:
        subset = eval_df[eval_df['System'] == sys]
        sys_summary.append([
            f'System {sys}',
            f"{subset['Time (seconds)'].mean():.1f}",
            f"{subset['Accuracy (0/1/2)'].mean():.2f}",
            f"{subset['Likert (1-5)'].mean():.2f}",
        ])
    make_table(doc,
        ['System', 'Avg Time (s)', 'Avg Accuracy', 'Avg Likert'],
        sys_summary
    )

    # Charts
    buf1 = create_time_by_system_task(eval_df)
    add_chart_image(doc, buf1,
        'Figure 4: Average task completion time by system and task')

    buf2 = create_accuracy_by_system_task(eval_df)
    add_chart_image(doc, buf2,
        'Figure 5: Average accuracy by system and task')

    buf3 = create_accuracy_heatmap(eval_df)
    add_chart_image(doc, buf3,
        'Figure 6: Accuracy heatmap (System x Task) — green indicates higher accuracy')

    buf4 = create_likert_by_system(eval_df)
    add_chart_image(doc, buf4,
        'Figure 7: Average ease-of-use rating by system')

    buf5 = create_ranking_chart(rank_df)
    add_chart_image(doc, buf5,
        'Figure 8: Post-study system preference rankings')

    add_heading(doc, '7.3 Analysis', level=2)

    add_body(doc, 'Task Completion Time', bold=True)
    add_body(doc,
        'System C was the fastest overall (mean 18.8s), followed by System A (21.0s) and '
        'System B (26.2s). System C\'s speed advantage was most pronounced on T2 and T3, where '
        'participants found the correct answer rapidly using legend-based filtering. System B '
        'was slowest particularly on T4 (36.9s average), where participants needed to adjust '
        'multiple dropdowns alongside the slider. This suggests that systems requiring fewer '
        'interaction steps for a given task lead to faster completion times.'
    )

    add_body(doc, 'Accuracy', bold=True)
    add_body(doc,
        'System A achieved the highest overall accuracy (1.68/2.0), followed by System B (1.56) '
        'and System C (1.52). System A excelled on T1 and T5, where its sorted bar chart with '
        'direct metric switching minimised interpretation errors. Notably, T2 (top 3 weather '
        'types by frequency) achieved the highest accuracy across all systems (1.80 average), '
        'indicating that frequency-counting tasks are well supported by bar chart idioms. T1 had '
        'the widest accuracy variation between systems: System A scored 1.60 while Systems B and '
        'C scored lower, likely because A\'s dual bar chart layout made seasonal/monthly comparison '
        'more straightforward.'
    )

    add_body(doc, 'Ease of Use (Likert)', bold=True)
    add_body(doc,
        'System A received the highest average Likert rating (4.20), followed by System C (4.08) '
        'and System B (3.88). Participant comments reveal that System A\'s straightforward bar '
        'chart layout felt intuitive, while System B\'s multiple dropdowns and heatmap colour '
        'scheme caused confusion. System C received mixed feedback: P3 and P5 rated it highly '
        'for its filtering flexibility, while P1 found the shape legend confusing ("scatter plot '
        'shapes made it more confusing").'
    )

    add_body(doc, 'Post-Study Rankings', bold=True)
    add_body(doc,
        'System B was ranked as the best system by 3 out of 5 participants, despite having the '
        'slowest average time and lowest Likert scores. This apparent contradiction suggests '
        'that participants valued the richness of temporal exploration and heatmap overview, even '
        'though it required more effort. System A was ranked worst by 3 participants, yet had the '
        'highest accuracy and Likert scores — indicating that while it was effective and easy, '
        'participants found it less engaging. System C received polarised rankings, being ranked '
        'best by P4 but worst by P1 and P5.'
    )

    add_body(doc, 'Key Findings', bold=True)
    add_bullet(doc,
        'No single system was uniformly "best" — each excelled in different dimensions '
        '(time, accuracy, preference).')
    add_bullet(doc,
        'Simpler interaction designs (System A) yielded higher accuracy and ease-of-use '
        'but lower engagement.')
    add_bullet(doc,
        'Richer visual encodings (System C\'s dual colour+shape) increased speed for some '
        'users but caused confusion for others.')
    add_bullet(doc,
        'Temporal overview tools (System B\'s heatmap) were valued for understanding despite '
        'slower task times.')

    # =====================================================================
    # PART 8: FUTURE WORK (max 400 words)
    # =====================================================================

    add_heading(doc, '8. Future Work', level=1)

    add_body(doc,
        'Based on our evaluation data, we identify the following evidence-based improvements:'
    )

    add_body(doc, '1. Reduce interaction complexity in System B', bold=True)
    add_body(doc,
        'System B had the slowest average completion time (26.2s) and lowest Likert score '
        '(3.88). Participant notes reveal that the need to adjust multiple dropdowns (Metric, '
        'Group By, Generalisation Level) created confusion. P1 "did not remember that the '
        'question was same" across systems because the interface required different mental models. '
        'We would consolidate the Metric and Group By dropdowns into a single task-oriented '
        'control, reducing the number of interaction steps required to answer each task.'
    )

    add_body(doc, '2. Simplify System C\'s visual encoding', bold=True)
    add_body(doc,
        'System C\'s dual encoding (colour for weather type, shape for season) caused confusion '
        'for P1 ("multiple colour in bar chart created the confusion," "scatter plot shapes made '
        'it more confusing") and P2 ("looking at scatter plot every time"). Our evaluation shows '
        'C had the lowest accuracy (1.52). We would remove the shape channel and instead use '
        'faceting by season (small multiples), which leverages the Gestalt principle of proximity '
        'to group seasonal data without overloading the shape channel (Munzner, 2014).'
    )

    add_body(doc, '3. Add a reset button to all systems', bold=True)
    add_body(doc,
        'Multiple participants (P1: "does not remember how to refresh the page"; P1: "did not '
        'refresh the page") struggled to reset filters between tasks. Currently, the only reset '
        'mechanism is refreshing the browser page. We would add an explicit "Reset All Filters" '
        'button within the interface, reducing the hidden knowledge required to operate the system.'
    )

    add_body(doc, '4. Improve task wording clarity', bold=True)
    add_body(doc,
        'Several participants (P1: "getting confuse in all tasks because of the wordings"; P2: '
        '"questions wording is confusing") struggled with task descriptions. While this is an '
        'evaluation design issue rather than a system design issue, we would revise the Tasks '
        'section in each system page to use simpler, step-by-step instructions with bold '
        'highlighting of key values and expected actions.'
    )

    # =====================================================================
    # REFERENCES
    # =====================================================================

    doc.add_page_break()
    add_heading(doc, 'References', level=1)

    refs = [
        'Mackinlay, J. (1986). Automating the design of graphical presentations of relational '
        'information. ACM Transactions on Graphics, 5(2), 110-141.',

        'Munzner, T. (2014). Visualization Analysis and Design. CRC Press.',

        'Shneiderman, B. (1996). The eyes have it: A task by data type taxonomy for information '
        'visualizations. Proceedings of the IEEE Symposium on Visual Languages, 336-343.',

        'Satyanarayan, A., Moritz, D., Wongsuphasawat, K., & Heer, J. (2017). Vega-Lite: A '
        'grammar of interactive graphics. IEEE Transactions on Visualization and Computer '
        'Graphics, 23(1), 341-350.',

        'Heer, J., & Shneiderman, B. (2012). Interactive dynamics for visual analysis. '
        'Communications of the ACM, 55(4), 45-54.',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Pt(36)
        p.paragraph_format.first_line_indent = Pt(-36)

    # =====================================================================
    # APPENDIX A: RAW EVALUATION DATA
    # =====================================================================

    doc.add_page_break()
    add_heading(doc, 'Appendix A: Raw Evaluation Data', level=1)

    add_body(doc,
        'Complete evaluation data from all 5 participants across all 3 systems and 5 tasks '
        '(75 observations).'
    )

    headers = ['P', 'Order', 'Sys', 'Task', 'Time(s)', 'Acc', 'Likert', 'Answer']
    rows = []
    for _, r in eval_df.iterrows():
        answer = str(r['Participant Answer'])[:40]
        rows.append([
            r['Participant'],
            str(r['System Order'])[:9],
            r['System'],
            r['Task'],
            f"{r['Time (seconds)']:.1f}",
            str(int(r['Accuracy (0/1/2)'])),
            str(int(r['Likert (1-5)'])),
            answer,
        ])
    make_table(doc, headers, rows)

    # Notes table
    add_heading(doc, 'Evaluator Notes', level=2)
    headers_notes = ['P', 'Sys', 'Task', 'Notes']
    rows_notes = []
    for _, r in eval_df.iterrows():
        notes = str(r.get('Notes', ''))
        if (notes and notes != 'nan' and notes.strip()
                and notes.strip().lower() != 'none'):
            rows_notes.append([
                r['Participant'],
                r['System'],
                r['Task'],
                notes[:80],
            ])
    if rows_notes:
        make_table(doc, headers_notes, rows_notes)

    # =====================================================================
    # APPENDIX B: POST-STUDY RANKINGS
    # =====================================================================

    add_heading(doc, 'Appendix B: Post-Study Rankings', level=1)

    headers_rank = ['Participant', 'Rank 1 (Best)', 'Rank 2', 'Rank 3 (Worst)', 'Comments']
    rows_rank = []
    for _, r in rank_df.iterrows():
        comment = str(r['Reason / Comments'])[:100]
        rows_rank.append([
            r['Participant'],
            r['Rank 1 (Best System)'],
            r['Rank 2'],
            r['Rank 3 (Worst System)'],
            comment,
        ])
    make_table(doc, headers_rank, rows_rank)

    # =====================================================================
    # APPENDIX C: TEAM CONTRIBUTION LOG
    # =====================================================================

    add_heading(doc, 'Appendix C: Team Contribution Log', level=1)

    add_body(doc, '[Please fill in team member names and contributions below]')

    contrib_headers = ['Team Member', 'Contribution', 'Estimated %']
    contrib_rows = [
        ['Member 1', '[e.g., System A implementation, data prep, report Part 1-2]', '20%'],
        ['Member 2', '[e.g., System B implementation, generalised selection, Part 4]', '20%'],
        ['Member 3', '[e.g., System C implementation, index page, report Part 3]', '20%'],
        ['Member 4', '[e.g., User evaluation, data collection, report Part 7-8]', '20%'],
        ['Member 5', '[e.g., Design comparison, demo video, report Part 6]', '20%'],
    ]
    make_table(doc, contrib_headers, contrib_rows)

    # =====================================================================
    # Save
    # =====================================================================

    os.makedirs(os.path.dirname(OUTPUT_DOCX), exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"Report saved to: {OUTPUT_DOCX}")


if __name__ == '__main__':
    build_report()
